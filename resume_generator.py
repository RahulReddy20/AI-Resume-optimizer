import os
import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from google import genai
from dotenv import load_dotenv
from google.genai import types
import re

# Load environment variables
load_dotenv()

# Configure the Gemini API client
client = genai.Client(api_key=os.getenv("GOOGLE_API_KEY"))


def generate_optimized_resume(resume_json, job_description, missing_skills, similarity_score):
    """
    Generate an optimized resume based on the job description using Gemini API

    Parameters:
        resume_text (str): Original resume text
        job_description (str): Job description text
        missing_skills (list): List of skills in the job description but not in the resume
        similarity_score (float): Similarity score between resume and job description

    Returns:
        dict: JSON structure of the optimized resume
    """
    # Prepare prompt for Gemini
    prompt = f"""
    You are a professional resume writer tasked with optimizing a resume to better match a job description.
    
    ORIGINAL RESUME in JSON format:
    {json.dumps(resume_json, indent=2)}
    
    JOB DESCRIPTION:
    {job_description}
    
    ANALYSIS:
    - Current match score: {similarity_score:.2f} out of 1.00
    - Missing skills/keywords: {', '.join(missing_skills) if isinstance(missing_skills, list) else str(missing_skills)}
    
    Please rewrite the resume to better match the job description while maintaining truthfulness.
    Focus on highlighting relevant experience, rewording skills, and restructuring content to showcase
    the candidate's fit for this specific position.
    
    YOUR OUTPUT MUST INCLUDE: 
    - Complete contact information
    - Education details
    - Skills formatted properly
    - Experience entries with descriptions
    - Project entries with descriptions if the Original Resume has projects
    
    Format the output as a JSON object in the same structure as the Original Resume.
    IMPORTANT: Ensure all property names and string values are enclosed in DOUBLE QUOTES, not single quotes.
    Follow the RFC 8259 JSON specification exactly.
    All strings MUST use proper escaping for special characters.
    
    Only include sections that are present in the original resume. Keep the content TRUTHFUL and based on the original resume.
    """

    try:
        # Generate content with Gemini
        response = client.models.generate_content(
            model="gemini-2.0-flash",
            contents=prompt,
            config=types.GenerateContentConfig(
                temperature=0.2,
                top_p=0.95,
                top_k=40,
                max_output_tokens=2048,
                response_mime_type="application/json"
            )
        )

        # Extract and parse JSON response
        try:
            # Try to parse the result directly
            resume_json = json.loads(response.text)
            # Validate the response has the expected structure
            if not isinstance(resume_json, dict):
                raise ValueError("Response is not a valid JSON object")

            print("Successfully parsed response JSON")
            return resume_json

        except json.JSONDecodeError as e:
            print(f"JSON parse error in resume generation: {str(e)}")

            # If direct parsing fails, try to extract JSON from text
            text = response.text

            # Find JSON content between curly braces
            start_idx = text.find('{')
            end_idx = text.rfind('}') + 1

            if start_idx >= 0 and end_idx > start_idx:
                json_str = text[start_idx:end_idx]
                try:
                    # Try to clean the JSON before parsing
                    cleaned_json = json_str

                    # Replace single quotes with double quotes (improved regex)
                    cleaned_json = re.sub(
                        r"(?<={|,)\s*'([^']+?)'(?=\s*:)", r'"\1"', cleaned_json)
                    cleaned_json = re.sub(
                        r":\s*'((?:[^'\\]|\\.)*)'/g", r':"\1"', cleaned_json)

                    # Fix escaped backslashes
                    cleaned_json = cleaned_json.replace("\\\\", "\\\\\\\\")

                    # Remove any code block markers
                    cleaned_json = re.sub(r'```json|```', '', cleaned_json)

                    # Fix unescaped newlines in string values
                    cleaned_json = re.sub(
                        r'"\s*\n\s*([^"])', r'" \1', cleaned_json)

                    # Handle multiline strings
                    cleaned_json = re.sub(r'"\s*\n\s*"', r'', cleaned_json)

                    resume_json = json.loads(cleaned_json)
                    print("Successfully extracted and cleaned JSON from response text")
                    return resume_json

                except json.JSONDecodeError as e2:
                    print(f"Secondary JSON parse error: {str(e2)}")

                    # Try a recursive approach with Gemini
                    for attempt in range(1, 4):  # Try up to 3 times
                        try:
                            fix_prompt = f"""
                            This JSON has syntax errors and cannot be parsed. Please fix it to be valid JSON with double quotes 
                            for all property names and string values. Fix all escaping and format issues:
                            
                            {cleaned_json if attempt == 1 else json_str}
                            
                            Return ONLY the fixed JSON with no additional text.
                            Make sure all strings are properly escaped with special attention to:
                            1. Ensure all backslashes are properly doubled when needed in strings
                            2. Fix any unescaped quotes inside string values
                            3. Fix unescaped newlines in string values
                            4. Ensure no trailing commas in arrays or objects
                            5. All keys and string values must use double quotes
                            """

                            fix_response = client.models.generate_content(
                                model="gemini-2.0-flash",
                                contents=fix_prompt,
                                config=types.GenerateContentConfig(
                                    temperature=0,
                                    response_mime_type="application/json"
                                )
                            )

                            # Try to parse the fixed JSON
                            fixed_text = fix_response.text.strip()

                            # Extract content between code blocks if present
                            if fixed_text.startswith("```") and "```" in fixed_text:
                                fixed_text = re.search(
                                    r'```(?:json)?\s*([\s\S]+?)\s*```', fixed_text).group(1)

                            # Try to find JSON structure
                            json_start = fixed_text.find('{')
                            json_end = fixed_text.rfind('}') + 1
                            if json_start >= 0 and json_end > json_start:
                                fixed_json = fixed_text[json_start:json_end]

                                try:
                                    resume_json = json.loads(fixed_json)
                                    print(
                                        f"Successfully fixed JSON with Gemini (attempt {attempt})")
                                    return resume_json
                                except json.JSONDecodeError:
                                    # Use this as input for next attempt
                                    json_str = fixed_json
                                    print(
                                        f"JSON fix attempt {attempt} failed, trying again")
                            else:
                                print(
                                    f"No JSON structure found in fix attempt {attempt}")

                        except Exception as e3:
                            print(f"Error in fix attempt {attempt}: {str(e3)}")
                            continue

                    # Last resort: Extract key-value pairs using regex
                    print("All Gemini repair attempts failed, trying regex extraction")
                    try:
                        pattern = r'"([^"]+)":\s*"([^"\\]*(?:\\.[^"\\]*)*)"'
                        matches = re.findall(pattern, json_str)
                        if matches:
                            # Create a minimal valid JSON structure with extracted key-values
                            extracted_json = {}
                            for key, value in matches:
                                value_cleaned = value.replace('\\"', '"')
                                parts = key.split('.')
                                if len(parts) == 1:
                                    extracted_json[key] = value_cleaned
                                else:
                                    # Handle nested keys like "contact_info.name"
                                    current = extracted_json
                                    for part in parts[:-1]:
                                        if part not in current:
                                            current[part] = {}
                                        current = current[part]
                                    current[parts[-1]] = value_cleaned

                            # Ensure minimal required structure
                            if extracted_json:
                                print("Created partial JSON from regex extraction")

                                # Add minimal required fields if missing
                                if "contact_info" not in extracted_json:
                                    extracted_json["contact_info"] = {
                                        "name": "Resume Owner"}
                                if "experience" not in extracted_json:
                                    extracted_json["experience"] = []
                                if "education" not in extracted_json:
                                    extracted_json["education"] = []

                                return extracted_json
                    except Exception as e4:
                        print(f"Regex extraction failed: {str(e4)}")
            else:
                raise ValueError(
                    "Could not extract valid JSON from response - no JSON structure found")

        # Create fallback minimal JSON with original resume data
        print("All JSON parsing attempts failed, creating fallback resume")
        fallback_resume = {
            "contact_info": {
                "name": resume_json.get("contact_info", {}).get("name", "Resume Owner"),
                "email": resume_json.get("contact_info", {}).get("email", ""),
                "phone": resume_json.get("contact_info", {}).get("phone", ""),
                "location": resume_json.get("contact_info", {}).get("location", "")
            },
            "summary": resume_json.get("summary", "Professional with relevant experience and skills."),
            "skills": resume_json.get("skills", {"technical_skills": []}),
            "experience": resume_json.get("experience", []),
            "education": resume_json.get("education", [])
        }

        print("Created minimal fallback resume structure from original data")
        return fallback_resume

    except Exception as e:
        print(f"Error generating optimized resume: {str(e)}")

        # Create fallback minimal JSON with original resume data if possible
        try:
            # Extract basic info for a minimal resume
            fallback_resume = {
                "contact_info": {
                    "name": resume_json.get("contact_info", {}).get("name", "Resume Owner"),
                    "email": resume_json.get("contact_info", {}).get("email", ""),
                    "phone": resume_json.get("contact_info", {}).get("phone", ""),
                    "location": resume_json.get("contact_info", {}).get("location", "")
                },
                "summary": resume_json.get("summary", "Professional with relevant experience and skills."),
                "skills": resume_json.get("skills", {"technical_skills": []}),
                "experience": resume_json.get("experience", []),
                "education": resume_json.get("education", [])
            }

            print("Created minimal fallback resume structure")
            return fallback_resume
        except:
            print("Could not create fallback resume")
            return None


def create_resume_docx(resume_json, output_path="optimized_resume.docx"):
    """
    Create a formatted Word document from resume JSON

    Parameters:
        resume_json (dict): Resume data in JSON format
        output_path (str): Output path for the Word document

    Returns:
        bool: Success status
    """
    try:
        # Create a new Document
        doc = Document()

        # Contact Information
        name = doc.add_paragraph()
        name_run = name.add_run(resume_json["contact_info"]["name"])
        name_run.bold = True
        name_run.font.size = Pt(16)
        name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        contact = doc.add_paragraph()
        contact_info = [
            resume_json["contact_info"].get("email", ""),
            resume_json["contact_info"].get("phone", ""),
            resume_json["contact_info"].get("location", "")
        ]
        if "linkedin" in resume_json["contact_info"]:
            contact_info.append(resume_json["contact_info"]["linkedin"])
        contact.add_run(" | ".join(filter(None, contact_info)))
        contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Summary
        doc.add_paragraph()
        summary_heading = doc.add_paragraph()
        summary_heading.add_run("SUMMARY").bold = True
        doc.add_paragraph(resume_json["summary"])

        # Skills
        doc.add_paragraph()
        skills_heading = doc.add_paragraph()
        skills_heading.add_run("SKILLS").bold = True
        skills_para = doc.add_paragraph()
        all_skills = []
        if isinstance(resume_json["skills"], list):
            all_skills = resume_json["skills"]
        elif isinstance(resume_json["skills"], dict):
            if "technical_skills" in resume_json["skills"]:
                all_skills.extend(resume_json["skills"]["technical_skills"])
            if "soft_skills" in resume_json["skills"]:
                all_skills.extend(resume_json["skills"]["soft_skills"])
            if "other_skills" in resume_json["skills"]:
                all_skills.extend(resume_json["skills"]["other_skills"])
        skills_para.add_run(", ".join(all_skills))

        # Experience
        doc.add_paragraph()
        exp_heading = doc.add_paragraph()
        exp_heading.add_run("EXPERIENCE").bold = True

        for job in resume_json["experience"]:
            job_title = doc.add_paragraph()
            title_company = f"{job['title']} - {job['company']}"
            job_title.add_run(title_company).bold = True
            dates = doc.add_paragraph()
            dates.add_run(job["dates"]).italic = True

            for bullet in job["description"]:
                bullet_point = doc.add_paragraph()
                bullet_point.style = 'List Bullet'
                bullet_point.add_run(bullet)

        # Education
        doc.add_paragraph()
        edu_heading = doc.add_paragraph()
        edu_heading.add_run("EDUCATION").bold = True

        for edu in resume_json["education"]:
            edu_degree = doc.add_paragraph()
            edu_degree.add_run(
                f"{edu['degree']} - {edu['institution']}").bold = True
            dates = doc.add_paragraph()
            dates.add_run(edu["dates"]).italic = True
            if "details" in edu and edu["details"]:
                details = doc.add_paragraph()
                details.add_run(edu["details"])

        # Projects (if available)
        if "projects" in resume_json and resume_json["projects"]:
            doc.add_paragraph()
            proj_heading = doc.add_paragraph()
            proj_heading.add_run("PROJECTS").bold = True

            for project in resume_json["projects"]:
                proj_name = doc.add_paragraph()
                proj_name.add_run(project["title"]).bold = True
                desc = doc.add_paragraph()
                desc.add_run(project["description"])

        # Certifications (if available)
        if "certifications" in resume_json and resume_json["certifications"]:
            doc.add_paragraph()
            cert_heading = doc.add_paragraph()
            cert_heading.add_run("CERTIFICATIONS").bold = True

            for cert in resume_json["certifications"]:
                cert_name = doc.add_paragraph()
                cert_name.add_run(
                    f"{cert['name']} - {cert['issuer']}").bold = True
                date = doc.add_paragraph()
                date.add_run(cert["date"]).italic = True

        # Activities (if available)
        if "activities" in resume_json and resume_json["activities"]:
            doc.add_paragraph()
            act_heading = doc.add_paragraph()
            act_heading.add_run("EXTRA-CURRICULAR ACTIVITIES").bold = True

            for activity in resume_json["activities"]:
                act_bullet = doc.add_paragraph()
                act_bullet.style = 'List Bullet'
                act_bullet.add_run(activity)

        # Leadership (if available)
        if "leadership" in resume_json and resume_json["leadership"]:
            doc.add_paragraph()
            lead_heading = doc.add_paragraph()
            lead_heading.add_run("LEADERSHIP").bold = True

            for lead_item in resume_json["leadership"]:
                lead_bullet = doc.add_paragraph()
                lead_bullet.style = 'List Bullet'
                lead_bullet.add_run(lead_item)

        # Save the document
        doc.save(output_path)
        return True

    except Exception as e:
        print(f"Error creating resume document: {str(e)}")
        return False


def create_resume_latex(resume_json, output_path="optimized_resume.tex"):
    """
    Create a LaTeX document from resume JSON using the template in latex_resume_format folder
    Uses Gemini API to directly generate LaTeX content.

    Parameters:
        resume_json (dict): Resume data in JSON format
        output_path (str): Output path for the LaTeX document

    Returns:
        bool: Success status
    """
    try:
        # Determine template file path
        template_dir = os.path.join(os.path.dirname(
            os.path.abspath(__file__)), "latex_resume_format")
        template_file = os.path.join(template_dir, "resume_faangpath.tex")

        # Check if template exists
        if not os.path.exists(template_file):
            print(f"Warning: Template file not found at {template_file}")
            print("Falling back to default LaTeX generation...")
            return _create_default_latex_resume(resume_json, output_path)

        # Read the template file
        with open(template_file, 'r', encoding='utf-8') as f:
            template_content = f.read()

        # Use Gemini to generate complete LaTeX document
        latex_content = analyze_and_map_template(template_content, resume_json)

        if latex_content:
            # Write the LaTeX content to the output file
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(latex_content)

            # Copy resume.cls to the output directory if it doesn't exist
            output_dir = os.path.dirname(output_path)
            if output_dir == "":
                output_dir = "."

            cls_source = os.path.join(template_dir, "resume.cls")
            cls_target = os.path.join(output_dir, "resume.cls")

            try:
                if os.path.exists(cls_source) and not os.path.exists(cls_target):
                    with open(cls_source, 'r', encoding='utf-8') as src:
                        with open(cls_target, 'w', encoding='utf-8') as dst:
                            dst.write(src.read())
                    print(f"Copied resume.cls file to {cls_target}")
                else:
                    print(f"Note: resume.cls already exists at {cls_target}")
            except Exception as e:
                print(f"Warning: Could not copy resume.cls file: {str(e)}")
                print(
                    f"You may need to manually copy resume.cls from {template_dir} to {output_dir}")

            print(
                f"✅ LaTeX resume file created at {output_path}")
            print("To generate PDF, run: pdflatex optimized_resume.tex")
            return True
        else:
            # Fallback to default LaTeX generation if Gemini API fails
            print(
                "Warning: Could not generate LaTeX with AI. Falling back to default generation.")
            return _create_default_latex_resume(resume_json, output_path)

    except Exception as e:
        print(f"Error creating LaTeX resume from template: {str(e)}")
        print(f"Debug info - JSON keys: {list(resume_json.keys())}")
        print("Falling back to default LaTeX generation...")
        return _create_default_latex_resume(resume_json, output_path)


def analyze_and_map_template(template_content, resume_json):
    """
    Use Gemini API to analyze LaTeX template and directly generate LaTeX content for the resume

    Parameters:
        template_content (str): LaTeX template content
        resume_json (dict): Resume data in JSON format

    Returns:
        str: Complete LaTeX document content or None if failed
    """
    try:
        # Extract sections from template for Gemini to analyze
        sections = re.findall(
            r'\\begin\{rSection\}\{([^}]*)\}', template_content)
        section_info = {}

        # Also check for commented-out sections like OBJECTIVE
        commented_sections = re.findall(
            r'%\s*\\begin\{rSection\}\{([^}]*)\}', template_content)

        # Create a prompt for Gemini to analyze template and generate LaTeX
        prompt = f"""
        You are an expert LaTeX document generator. I need you to create a well-formatted resume in LaTeX using a specific template structure.
        
        TEMPLATE SECTIONS:
        {', '.join(sections)}
        
        COMMENTED SECTIONS (can be uncommented):
        {', '.join(commented_sections)}
        
        JSON RESUME DATA:
        {json.dumps(resume_json, indent=2)}
        
        CURRENT TEMPLATE STRUCTURE:
        {template_content}
        
        I need you to generate a complete LaTeX document using the existing template structure. 
        Please follow these requirements:
        1. Keep the document class and package definitions exactly as they are
        2. Use the \\name{{}} command to set the person's name from contact_info
        3. Use the \\address{{}} commands for contact information like phone, email, location, etc.
        4. For each section in the template (rSection), fill in appropriate content from the JSON data
        5. Ensure proper LaTeX escaping for special characters (_, %, $, #, &, etc.)
        6. Format each section according to the template's existing style and spacing
        7. Maintain the overall document structure with \\begin{{document}} and \\end{{document}}
       
        8. IMPORTANT: For the ADDRESS section, modify the address command to use wrapping, either by:
           - See if the address is too long, if it is then break the address into multiple address blocks
           - Using shorter lines that don't exceed page width\
           - Breaking long text with explicit line breaks (\\\\)
           - For URLs, use shorter display text: \\href{{https://www.linkedin.com/in/username}}{{LinkedIn}}
        
        Return ONLY the complete LaTeX document with no additional text before or after.
        """

        # Generate content with Gemini
        response = client.models.generate_content(
            model="gemini-2.0-flash",
            contents=prompt,
            config=types.GenerateContentConfig(
                temperature=0.0,
                top_p=0.5,
                top_k=40,
                max_output_tokens=8192,
                response_mime_type="text/plain"
            )
        )

        # Process the response as raw text
        latex_content = response.text.strip()

        # Remove any markdown code block markers more thoroughly
        if "```" in latex_content:
            # First try the standard code block extraction
            match = re.search(
                r'```(?:latex)?[\s\r\n]*([\s\S]+?)[\s\r\n]*```', latex_content)
            if match:
                latex_content = match.group(1).strip()
                print("Successfully extracted LaTeX content from code block")
            else:
                # If regex match fails, try simple removal
                latex_content = latex_content.replace(
                    "```latex", "").replace("```", "").strip()
                print("Removed code block markers from LaTeX content")

        # Check if content starts with LaTeX document indicator
        if "\\documentclass" in latex_content:
            print("Successfully generated LaTeX content directly")
            return latex_content
        else:
            print(
                "Warning: Generated content may not be valid LaTeX. Attempting to use it anyway.")
            return latex_content

    except Exception as e:
        print(f"Error generating LaTeX with Gemini API: {str(e)}")
        return None


def create_fallback_mapping(resume_json, sections):
    """
    Create a fallback mapping when Gemini API JSON parsing fails

    Parameters:
        resume_json (dict): Resume data in JSON format
        sections (list): List of section names from template

    Returns:
        dict: Basic mapping of sections to LaTeX content
    """
    mapping = {}

    # Add contact info
    if "contact_info" in resume_json:
        contact = resume_json["contact_info"]
        name = escape_latex(str(contact.get("name", "")))
        mapping["name"] = name

        # Create address blocks with proper wrapping
        phone = escape_latex(str(contact.get("phone", "")))
        location = escape_latex(str(contact.get("location", "")))
        # Simple line breaks for address
        address1 = f"{phone} \\\\ {location}"
        mapping["address1"] = address1

        email = escape_latex(str(contact.get("email", "")))
        linkedin = escape_latex(str(contact.get("linkedin", "")))
        github = escape_latex(str(contact.get("github", "")))
        website = escape_latex(str(contact.get("website", "")))

        # Format links with shorter display text and explicit line breaks
        address2_parts = []
        if email:
            address2_parts.append(f"\\href{{mailto:{email}}}{{{email}}}")
        if linkedin:
            # Use shorter display text for LinkedIn
            if "linkedin.com" in linkedin:
                address2_parts.append(f"\\href{{{linkedin}}}{{LinkedIn}}")
            else:
                address2_parts.append(linkedin)
        if github:
            # Use shorter display text for GitHub
            if "github.com" in github:
                address2_parts.append(f"\\href{{{github}}}{{GitHub}}")
            else:
                address2_parts.append(github)
        if website:
            # Use shorter display text for website
            if len(website) > 30:
                domain = website.replace(
                    "https://", "").replace("http://", "").split("/")[0]
                address2_parts.append(f"\\href{{{website}}}{{{domain}}}")
            else:
                address2_parts.append(f"\\href{{{website}}}{{{website}}}")

        # Join with explicit line breaks
        address2 = " \\\\ ".join(address2_parts)
        mapping["address2"] = address2

    # Map common sections
    for section in sections:
        section_lower = section.lower()

        # Education section
        if section_lower == "education" and "education" in resume_json:
            content = "\n\n"
            for edu in resume_json["education"]:
                degree = escape_latex(str(edu.get("degree", "")))
                institution = escape_latex(str(edu.get("institution", "")))
                dates = escape_latex(str(edu.get("dates", "")))
                details = escape_latex(str(edu.get("details", "")))

                content += f"{{\\bf {degree}}}, {institution} \\hfill {{{dates}}}\\\\\n"
                if details:
                    content += f"\\textbf{{Relevant Coursework:}} {details}\n\n"
                else:
                    content += "\n"

            mapping[section] = content

        # Skills section
        elif section_lower == "skills" and "skills" in resume_json:
            content = "\n\n\\begin{tabular}{ @{} >{\\bfseries}l @{\\hspace{4ex}} p{13cm} }"

            if isinstance(resume_json["skills"], list):
                skills_list = [escape_latex(str(skill))
                               for skill in resume_json["skills"]]
                skills_str = ", ".join(skills_list)
                content += f"Skills & {skills_str} \\\\"
            elif isinstance(resume_json["skills"], dict):
                for skill_category, skills in resume_json["skills"].items():
                    if skills and isinstance(skills, list):
                        category_display = escape_latex(
                            skill_category.replace("_", " ").title())
                        skills_str = ", ".join(
                            [escape_latex(str(skill)) for skill in skills])
                        content += f"{category_display} & {skills_str} \\\\"

            content += "\\end{tabular}\\\\\n"
            mapping[section] = content

        # Experience section
        elif (section_lower == "experience" or section_lower == "work experience") and "experience" in resume_json:
            content = "\n\n"

            for job in resume_json["experience"]:
                title = escape_latex(str(job.get("title", "")))
                company = escape_latex(str(job.get("company", "")))
                dates = escape_latex(str(job.get("dates", "")))
                location = escape_latex(str(job.get("location", "")))
                description = job.get("description", [])

                if not isinstance(description, list):
                    description = [str(description)]

                content += f"\\textbf{{{title}}} \\hfill {dates}\\\\\n"
                content += f"{company} \\hfill \\textit{{{location}}}\n"
                content += "\\begin{itemize}\n    \\itemsep -3pt {} \n"

                for bullet in description:
                    bullet_text = escape_latex(str(bullet))
                    content += f"     \\item {bullet_text}\n"

                content += "\\end{itemize}\n\n"

            mapping[section] = content

    # Add generic mapping for other sections if present in resume_json
    if "summary" in resume_json or "objective" in resume_json:
        summary = escape_latex(
            str(resume_json.get("summary", resume_json.get("objective", ""))))
        mapping["OBJECTIVE"] = f"\n\n{{{summary}}}\n\n"

    return mapping


def apply_gemini_mapping(template_content, gemini_mapping):
    """
    Apply the Gemini-generated section mappings to the LaTeX template

    Parameters:
        template_content (str): Original LaTeX template content
        gemini_mapping (dict): Mapping of sections to LaTeX content

    Returns:
        str: Modified template content with mapped sections
    """
    try:
        # Process each section in the mapping
        for section_name, section_content in gemini_mapping.items():
            # Check if section exists in template
            section_pattern = rf'\\begin{{rSection}}{{{section_name}}}(.*?)\\end{{rSection}}'
            section_match = re.search(
                section_pattern, template_content, re.DOTALL)

            if section_match:
                # Replace existing section
                replacement = f"\\begin{{rSection}}{{{section_name}}}\n{section_content}\n\\end{{rSection}}"
                template_content = template_content.replace(
                    section_match.group(0), replacement)
            else:
                # Check if it's a commented section that can be uncommented
                commented_pattern = rf'%\s*\\begin{{rSection}}{{{section_name}}}(.*?)%\s*\\end{{rSection}}'
                commented_match = re.search(
                    commented_pattern, template_content, re.DOTALL)

                if commented_match:
                    # Uncomment and replace section
                    commented_section = commented_match.group(0)
                    uncommented_section = commented_section.replace('% ', '')
                    replacement = f"\\begin{{rSection}}{{{section_name}}}\n{section_content}\n\\end{{rSection}}"
                    template_content = template_content.replace(
                        uncommented_section, replacement)
                else:
                    # Add new section at end of document
                    end_document_pos = template_content.find("\\end{document}")
                    if end_document_pos > 0 and section_content.strip():
                        new_section = f"\\begin{{rSection}}{{{section_name}}}\n{section_content}\n\\end{{rSection}}\n\n"
                        template_content = template_content[:end_document_pos] + \
                            new_section + template_content[end_document_pos:]

        # Handle contact information and name separately since they're not in rSections
        if "name" in gemini_mapping:
            name_pattern = r'\\name\{[^}]*\}'
            template_content = re.sub(
                name_pattern, f'\\name{{{gemini_mapping["name"]}}}', template_content)

        if "address1" in gemini_mapping and "address2" in gemini_mapping:
            address_patterns = list(re.finditer(
                r'\\address\{[^}]*\}', template_content))
            if len(address_patterns) >= 2:
                template_content = re.sub(re.escape(address_patterns[0].group(0)),
                                          f'\\address{{{gemini_mapping["address1"]}}}',
                                          template_content, count=1)
                template_content = re.sub(re.escape(address_patterns[1].group(0)),
                                          f'\\address{{{gemini_mapping["address2"]}}}',
                                          template_content, count=1)
            elif len(address_patterns) == 1:
                template_content = re.sub(re.escape(address_patterns[0].group(0)),
                                          f'\\address{{{gemini_mapping["address1"]}}}\\address{{{gemini_mapping["address2"]}}}',
                                          template_content, count=1)

        return template_content

    except Exception as e:
        print(f"Error applying AI-generated mappings to template: {str(e)}")
        return template_content


def escape_latex(text):
    """
    Escape LaTeX special characters in text

    Parameters:
        text (str): Text to escape

    Returns:
        str: Text with LaTeX special characters escaped
    """
    if not isinstance(text, str):
        text = str(text)

    # Define LaTeX special characters that need escaping
    special_chars = {
        '&': '\\&',
        '%': '\\%',
        '$': '\\$',
        '#': '\\#',
        '_': '\\_',
        '{': '\\{',
        '}': '\\}',
        '~': '\\textasciitilde{}',
        '^': '\\textasciicircum{}',
        '\\': '\\textbackslash{}',
        '<': '\\textless{}',
        '>': '\\textgreater{}'
    }

    # Replace each special character with its escaped version
    for char, replacement in special_chars.items():
        text = text.replace(char, replacement)

    return text

# Default LaTeX generation function for fallback


def _create_default_latex_resume(resume_json, output_path="optimized_resume.tex"):
    """
    Create a basic LaTeX document from resume JSON (fallback implementation)

    Parameters:
        resume_json (dict): Resume data in JSON format
        output_path (str): Output path for the LaTeX document

    Returns:
        bool: Success status
    """
    try:
        # Start building the LaTeX document
        latex_content = []

        # Document class and packages
        latex_content.append("\\documentclass{resume}")
        latex_content.append(
            "\\usepackage[left=0.4 in,top=0.4in,right=0.4 in,bottom=0.4in]{geometry}")
        latex_content.append(
            "\\newcommand{\\tab}[1]{\\hspace{.2667\\textwidth}\\rlap{#1}}")
        latex_content.append(
            "\\newcommand{\\itab}[1]{\\hspace{0em}\\rlap{#1}}")

        # Contact information
        name = escape_latex(
            str(resume_json["contact_info"].get("name", "Firstname Lastname")))
        latex_content.append(f"\\name{{{name}}}")

        # Contact info - first address block (phone, location)
        phone = escape_latex(str(resume_json["contact_info"].get("phone", "")))
        location = escape_latex(
            str(resume_json["contact_info"].get("location", "")))
        contact_line1 = f"{phone} \\\\ {location}"
        latex_content.append(f"\\address{{{contact_line1}}}")

        # Contact info - second address block (email, linkedin, website, github)
        email = escape_latex(str(resume_json["contact_info"].get("email", "")))
        linkedin = escape_latex(
            str(resume_json["contact_info"].get("linkedin", "")))
        website = escape_latex(
            str(resume_json["contact_info"].get("website", "")))
        github = escape_latex(
            str(resume_json["contact_info"].get("github", "")))

        # Format each contact element with shorter display text
        contact_parts = []
        if email:
            contact_parts.append(f"\\href{{mailto:{email}}}{{{email}}}")

        if linkedin:
            # Use shorter display text for LinkedIn
            if "linkedin.com" in linkedin:
                contact_parts.append(f"\\href{{{linkedin}}}{{LinkedIn}}")
            else:
                contact_parts.append(linkedin)

        if github:
            # Use shorter display text for GitHub
            if "github.com" in github:
                contact_parts.append(f"\\href{{{github}}}{{GitHub}}")
            else:
                contact_parts.append(github)

        if website:
            # Use shorter display text for website if URL is long
            if len(website) > 30:
                domain = website.replace(
                    "https://", "").replace("http://", "").split("/")[0]
                contact_parts.append(f"\\href{{{website}}}{{{domain}}}")
            else:
                contact_parts.append(f"\\href{{{website}}}{{{website}}}")

        # Join with explicit line breaks for better spacing
        contact_line2 = " \\\\ ".join(
            contact_parts) if contact_parts else "example@email.com"
        latex_content.append(f"\\address{{{contact_line2}}}")

        # Begin document
        latex_content.append("\\begin{document}")

        # Objective/Summary
        summary_text = escape_latex(str(resume_json.get("summary", resume_json.get(
            "objective", "Professional seeking opportunities to apply skills and experience."))))
        latex_content.append("\\begin{rSection}{OBJECTIVE}")
        latex_content.append("")
        latex_content.append(f"{{{summary_text}}}")
        latex_content.append("")
        latex_content.append("\\end{rSection}")

        # Education
        if "education" in resume_json and resume_json["education"]:
            latex_content.append("\\begin{rSection}{Education}")
            latex_content.append("")

            for edu in resume_json["education"]:
                degree = escape_latex(str(edu.get("degree", "")))
                institution = escape_latex(str(edu.get("institution", "")))
                dates = escape_latex(str(edu.get("dates", "")))
                details = escape_latex(str(edu.get("details", "")))

                latex_content.append(
                    f"{{\\bf {degree}}}, {institution} \\hfill {{{dates}}}")
                if details:
                    latex_content.append(f"Relevant Coursework: {details}")
                latex_content.append("")

            latex_content.append("\\end{rSection}")

        # Skills
        if "skills" in resume_json:
            latex_content.append("\\begin{rSection}{SKILLS}")
            latex_content.append("")
            latex_content.append(
                "\\begin{tabular}{ @{} >{\\bfseries}l @{\\hspace{4ex}} p{13cm} }")

            if isinstance(resume_json["skills"], list):
                skills_list = [escape_latex(str(skill))
                               for skill in resume_json["skills"]]
                skills_str = ", ".join(skills_list)
                latex_content.append(f"Skills & {skills_str} \\\\")
            elif isinstance(resume_json["skills"], dict):
                if "technical_skills" in resume_json["skills"]:
                    tech_skills_list = [escape_latex(
                        str(skill)) for skill in resume_json["skills"]["technical_skills"]]
                    tech_skills = ", ".join(tech_skills_list)
                    latex_content.append(
                        f"Technical Skills & {tech_skills} \\\\")
                if "soft_skills" in resume_json["skills"]:
                    soft_skills_list = [escape_latex(
                        str(skill)) for skill in resume_json["skills"]["soft_skills"]]
                    soft_skills = ", ".join(soft_skills_list)
                    latex_content.append(f"Soft Skills & {soft_skills} \\\\")
                if "other_skills" in resume_json["skills"]:
                    other_skills_list = [escape_latex(
                        str(skill)) for skill in resume_json["skills"]["other_skills"]]
                    other_skills = ", ".join(other_skills_list)
                    latex_content.append(f"Other Skills & {other_skills} \\\\")

            latex_content.append("\\end{tabular}\\\\")
            latex_content.append("\\end{rSection}")

        # Experience
        if "experience" in resume_json and resume_json["experience"]:
            latex_content.append("\\begin{rSection}{EXPERIENCE}")
            latex_content.append("")

            for job in resume_json["experience"]:
                title = escape_latex(str(job.get("title", "")))
                company = escape_latex(str(job.get("company", "")))
                dates = escape_latex(str(job.get("dates", "")))
                location = escape_latex(str(job.get("location", "")))
                description = job.get("description", [])

                latex_content.append(
                    f"\\textbf{{{title}}} \\hfill {dates}\\\\")
                latex_content.append(
                    f"{company} \\hfill \\textit{{{location}}}")
                latex_content.append("\\begin{itemize}")
                latex_content.append("    \\itemsep -3pt {} ")

                for bullet in description:
                    # Safely handle LaTeX special characters
                    safe_bullet = escape_latex(str(bullet))
                    latex_content.append(f"     \\item {safe_bullet}")

                latex_content.append("\\end{itemize}")
                latex_content.append("")

            latex_content.append("\\end{rSection}")

        # Projects
        if "projects" in resume_json and resume_json["projects"]:
            latex_content.append("\\begin{rSection}{PROJECTS}")
            latex_content.append("\\vspace{-1.25em}")

            for project in resume_json["projects"]:
                if isinstance(project, dict):
                    title = escape_latex(str(project.get("title", "")))
                    description = project.get("description", "")
                    technologies = escape_latex(
                        str(project.get("technologies", "")))
                    url = escape_latex(str(project.get("url", "")))

                    # Format project entry
                    if url:
                        title_text = f"\\textbf{{{title}}} \\href{{{url}}}{{(Link)}}"
                    else:
                        title_text = f"\\textbf{{{title}}}"

                    if technologies:
                        title_text += f" - {technologies}"

                    latex_content.append(f"\\item {title_text}")

                    # Handle description based on its type
                    if description:
                        if isinstance(description, list):
                            latex_content.append("\\begin{itemize}")
                            for bullet in description:
                                safe_bullet = escape_latex(str(bullet))
                                latex_content.append(
                                    f"    \\item {safe_bullet}")
                            latex_content.append("\\end{itemize}")
                        else:
                            safe_description = escape_latex(str(description))
                            latex_content.append(f"    {safe_description}")
                elif isinstance(project, str):
                    safe_project = escape_latex(str(project))
                    latex_content.append(f"\\item {safe_project}")

            latex_content.append("\\end{rSection}")

        # Activities (if available)
        if "activities" in resume_json and resume_json["activities"]:
            latex_content.append(
                "\\begin{rSection}{Extra-Curricular Activities}")
            latex_content.append("\\begin{itemize}")

            for activity in resume_json["activities"]:
                # Safely handle LaTeX special characters
                safe_activity = escape_latex(str(activity))
                latex_content.append(f"    \\item {safe_activity}")

            latex_content.append("\\end{itemize}")
            latex_content.append("\\end{rSection}")

        # Leadership (if available)
        if "leadership" in resume_json and resume_json["leadership"]:
            latex_content.append("\\begin{rSection}{Leadership}")
            latex_content.append("\\begin{itemize}")

            for lead_item in resume_json["leadership"]:
                # Safely handle LaTeX special characters
                safe_lead_item = escape_latex(str(lead_item))
                latex_content.append(f"    \\item {safe_lead_item}")

            latex_content.append("\\end{itemize}")
            latex_content.append("\\end{rSection}")

        # End document
        latex_content.append("\\end{document}")

        # Write to file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("\n".join(latex_content))

        # Also copy the resume.cls file to the output directory if it doesn't exist
        output_dir = os.path.dirname(output_path)
        if output_dir == "":
            output_dir = "."

        template_dir = os.path.join(os.path.dirname(
            os.path.abspath(__file__)), "latex_resume_format")
        cls_source = os.path.join(template_dir, "resume.cls")
        cls_target = os.path.join(output_dir, "resume.cls")

        try:
            if os.path.exists(cls_source) and not os.path.exists(cls_target):
                with open(cls_source, 'r', encoding='utf-8') as src:
                    with open(cls_target, 'w', encoding='utf-8') as dst:
                        dst.write(src.read())
        except Exception as e:
            print(f"Warning: Could not copy resume.cls file: {str(e)}")
            print(
                f"You may need to manually copy resume.cls from the latex_resume_format directory to {output_dir}")

        print(f"✅ LaTeX resume file created at {output_path}")
        print("To generate PDF, run: pdflatex optimized_resume.tex")
        return True

    except Exception as e:
        print(f"Error creating LaTeX resume: {str(e)}")
        print(f"Debug info - JSON keys: {list(resume_json.keys())}")
        return False
